from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCREEN_DIR = ROOT / "complete-project-screenshots"
OUT_PATH = Path(r"C:\Users\DELL\Desktop\Graduation Project Frontend Final.docx")


def set_run_font(run, name: str, size: int, bold: bool = False, color: tuple[int, int, int] | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 18, (0, 0, 0)),
        ("Heading 2", 15, (0, 0, 0)),
        ("Heading 3", 13, (67, 67, 67)),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_run_font(run, "Arial", 24, False, (0, 0, 0))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(14)
    run2 = p2.add_run(subtitle)
    set_run_font(run2, "Arial", 11, False, (85, 85, 85))


def add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_figure(doc: Document, filename: str, caption: str, width: float = 6.0) -> None:
    path = SCREEN_DIR / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, "Arial", 10, False, (85, 85, 85))


def main() -> None:
    doc = Document()
    configure_document(doc)

    add_title(
        doc,
        "CMO.AI Frontend Documentation",
        "Graduation Project Chapter Draft Based on the Implemented Frontend",
    )

    doc.add_heading("1. Frontend Overview", level=1)
    add_body(
        doc,
        "The frontend of CMO.AI is a single-page web application designed to guide the user from the first introduction to the platform into a protected AI-assisted marketing workspace. "
        "Its role is not limited to presentation; it also manages user authentication, subscription selection, feature exploration, and the operational dashboard used for campaign planning and content production."
    )
    add_body(
        doc,
        "From an architectural perspective, the frontend is divided into two major layers. The first layer contains the public pages such as the welcome page, landing page, pricing page, payment page, and authentication pages. "
        "The second layer is the authenticated dashboard, which provides access to the main product modules including orchestration, market planning, brand coaching, calendar planning, text generation, image generation, video generation, and performance analytics."
    )
    add_figure(doc, "01-welcome.png", "Figure 3.1. Welcome Page")
    add_figure(doc, "02-landing.png", "Figure 3.2. Landing Page")

    doc.add_heading("2. Frontend Technology Stack and Justification", level=1)
    add_body(
        doc,
        "The frontend is implemented using React, TypeScript, Vite, Tailwind CSS, shadcn/ui, Radix UI primitives, and Lucide icons. This stack was selected to support modularity, fast iteration, maintainability, and a modern user interface suitable for an AI-driven platform."
    )
    doc.add_heading("2.1 Build Tool: Vite", level=2)
    add_body(
        doc,
        "Vite is used as the frontend build tool because it provides a fast development server, efficient module handling, and optimized production builds. Compared to older alternatives such as Create React App, it reduces startup time and improves the development workflow."
    )
    doc.add_heading("2.2 UI Framework: React", level=2)
    add_body(
        doc,
        "React was chosen because of its component-based architecture. This approach makes it possible to build reusable interface units such as navigation elements, dashboard panels, cards, forms, and dialogs. "
        "For a system like CMO.AI, where many pages share layout patterns but differ in behavior, component reusability is essential."
    )
    doc.add_heading("2.3 Language: TypeScript", level=2)
    add_body(
        doc,
        "TypeScript improves code reliability through static typing. Since the application exchanges structured data with a FastAPI backend, type safety helps reduce integration errors, improves maintainability, and makes the codebase easier to scale."
    )
    doc.add_heading("2.4 Styling System: Tailwind CSS", level=2)
    add_body(
        doc,
        "Tailwind CSS is used for utility-first styling. It allows the interface to be built quickly while preserving visual consistency. The design system of CMO.AI uses dark gradients, glowing highlights, rounded panels, and responsive spacing to present a professional but modern dashboard experience."
    )
    doc.add_heading("2.5 Component System: shadcn/ui and Radix", level=2)
    add_body(
        doc,
        "shadcn/ui components provide reusable building blocks such as inputs, buttons, dialogs, tables, labels, and sheets. These components are built on accessible Radix primitives, which improves usability while keeping the design customizable."
    )
    doc.add_heading("2.6 Iconography: Lucide React", level=2)
    add_body(
        doc,
        "Lucide React is used as the icon library across the application. It provides lightweight vector icons that visually clarify actions, modules, and metrics without adding heavy visual overhead."
    )

    doc.add_heading("3. Frontend Architecture and Routing", level=1)
    add_body(
        doc,
        "The application starts from the React entry point in main.tsx, where the app is mounted inside BrowserRouter. Routing is handled by React Router. "
        "The main route configuration includes the public pages, authentication pages, dynamic feature pages, and the protected dashboard route."
    )
    add_body(
        doc,
        "The implemented frontend routes include the following paths: the welcome page (/), landing page (/landing), login page (/login), registration page (/register), forgot-password page (/forgot-password), OTP verification page (/verify-otp), reset-password page (/reset-password), pricing page (/pricing), payment page (/payment), dynamic feature pages (/features/:id), and the protected dashboard page (/dashboard)."
    )
    add_body(
        doc,
        "Authentication state is managed using AuthContext, while campaign and brand workspace state is managed through a dedicated CampaignProvider. This separation keeps route protection and workspace loading organized."
    )

    doc.add_heading("4. Public Pages", level=1)
    doc.add_heading("4.1 Welcome Page", level=2)
    add_body(
        doc,
        "The welcome page is the first screen shown to the user. Its purpose is to create an immediate first impression, communicate the platform identity, and guide the user either to explore the platform or to sign in. "
        "The page uses a full-screen hero layout, branded background imagery, clear typography, and two primary call-to-action buttons."
    )
    add_bullets(
        doc,
        [
            "Headline presenting the platform as an AI-powered marketing system.",
            "Branded logo and visual identity.",
            "Background image with overlay to improve readability.",
            "Primary actions to explore the platform or sign in.",
        ],
    )
    add_figure(doc, "01-welcome.png", "Figure 3.3. Welcome Page Interface")

    doc.add_heading("4.2 Landing Page", level=2)
    add_body(
        doc,
        "The landing page acts as the main marketing page of the application. It provides a broader explanation of the product, introduces its features, and guides users toward sign-up, pricing, or dashboard access depending on their state."
    )
    add_body(
        doc,
        "Its implementation is section-based and composed from reusable parts such as the navigation bar, hero section, features section, process overview, call-to-action section, and footer. "
        "This structure makes the page easy to maintain and easy to extend."
    )
    add_figure(doc, "02-landing.png", "Figure 3.4. Landing Page Structure")

    doc.add_heading("4.3 Pricing and Payment Pages", level=2)
    add_body(
        doc,
        "The pricing page displays the platform subscription plans and helps users compare available options. It supports the product’s conversion flow by communicating value, differentiating plans, and answering frequently asked questions."
    )
    add_body(
        doc,
        "The payment page is implemented as a conditional interface. When no plan is selected, the page shows a fallback message. When a plan is selected, the interface adapts to the selected plan and displays a checkout form. "
        "Free plans are activated directly, while paid plans require card-related form fields."
    )
    add_figure(doc, "03-pricing.png", "Figure 3.5. Pricing Page")
    add_figure(doc, "04-payment-no-plan.png", "Figure 3.6. Payment Page without Selected Plan")
    add_figure(doc, "05-payment-free.png", "Figure 3.7. Payment Page for Free Plan")
    add_figure(doc, "06-payment-pro.png", "Figure 3.8. Payment Page for Pro Plan")

    doc.add_heading("5. Authentication and Password Recovery", level=1)
    add_body(
        doc,
        "The authentication flow is a core part of the frontend because the main dashboard is protected. The system includes registration, login, password recovery, OTP verification, and password reset screens, all built on a shared authentication layout for visual consistency."
    )
    add_body(
        doc,
        "It is important to note that the implemented frontend does not contain a standalone Gmail page. Instead, it provides an email-based OTP recovery flow. "
        "The email address is entered in the forgot-password screen, the one-time code is entered in the OTP verification screen, and the new password is entered in the reset-password screen."
    )
    add_figure(doc, "07-login.png", "Figure 3.9. Login Page")
    add_figure(doc, "08-register.png", "Figure 3.10. Registration Page")
    add_figure(doc, "09-forgot-password.png", "Figure 3.11. Forgot Password Page")
    add_figure(doc, "10-verify-otp.png", "Figure 3.12. OTP Verification Page")
    add_figure(doc, "11-reset-password.png", "Figure 3.13. Reset Password Page")

    doc.add_heading("6. Feature Detail Pages", level=1)
    add_body(
        doc,
        "The landing page includes a feature section with \"Learn More\" actions that open dedicated feature pages. These pages give deeper explanations of each major capability before the user enters the dashboard. "
        "All feature detail pages are generated from a structured data source, which makes the implementation data-driven instead of hardcoded page by page."
    )
    add_bullets(
        doc,
        [
            "Brand Coaching",
            "Market Planning",
            "Smart Calendar",
            "Content Generation",
            "Performance Analytics",
            "Campaign Management",
        ],
    )
    add_body(
        doc,
        "Each feature page contains a hero section, an explanation of why the feature matters, a list of user actions supported by the feature, a practical use case, a description of how the feature connects to the larger platform workflow, and a call-to-action that encourages the user to continue."
    )
    add_figure(doc, "12-feature-brand-coaching.png", "Figure 3.14. Brand Coaching Feature Page")
    add_figure(doc, "13-feature-market-planning.png", "Figure 3.15. Market Planning Feature Page")
    add_figure(doc, "14-feature-smart-calendar.png", "Figure 3.16. Smart Calendar Feature Page")
    add_figure(doc, "15-feature-content-generation.png", "Figure 3.17. Content Generation Feature Page")
    add_figure(doc, "16-feature-analytics.png", "Figure 3.18. Performance Analytics Feature Page")
    add_figure(doc, "17-feature-campaign-management.png", "Figure 3.19. Campaign Management Feature Page")

    doc.add_heading("7. Dashboard Workspace", level=1)
    add_body(
        doc,
        "The dashboard is the operational core of the frontend. Unlike the public pages, it is not a simple informational interface. It acts as a unified workspace where the user manages campaigns, navigates between product modules, reviews metrics, and triggers AI-assisted actions."
    )
    add_body(
        doc,
        "The dashboard is protected through route-based authentication and is rendered only after the user session is validated. "
        "Inside the dashboard, the main content area changes dynamically based on the active workspace selected by the user."
    )

    doc.add_heading("7.1 Orchestrator", level=2)
    add_body(
        doc,
        "The Orchestrator serves as the command center of the dashboard. It summarizes campaign readiness, displays audience and launch context, and offers quick actions that guide the user to the next logical workflow step."
    )
    add_figure(doc, "18-dashboard-orchestrator.png", "Figure 3.20. Dashboard Orchestrator View")

    doc.add_heading("7.2 Market Planner", level=2)
    add_body(
        doc,
        "The Market Planner collects campaign inputs such as brand identity, target audience, industry, budget, product or service description, and selected platforms. "
        "Based on this information, it presents a structured plan including content pillars, posting recommendations, and execution steps."
    )
    add_figure(doc, "19-dashboard-market-planner.png", "Figure 3.21. Dashboard Market Planner View")

    doc.add_heading("7.3 Brand Coaching", level=2)
    add_body(
        doc,
        "The Brand Coaching workspace focuses on positioning, voice, and audience fit. It helps users clarify how the brand should communicate before content is created across campaigns."
    )
    add_figure(doc, "20-dashboard-brand-coaching.png", "Figure 3.22. Dashboard Brand Coaching View")

    doc.add_heading("7.4 Market Calendar", level=2)
    add_body(
        doc,
        "The Market Calendar supports planning and scheduling. It displays campaign timing logic and offers actions such as generating the next period of content, balancing channels, and identifying calendar gaps."
    )
    add_figure(doc, "21-dashboard-market-calendar.png", "Figure 3.23. Dashboard Market Calendar View")

    doc.add_heading("7.5 Text Generation", level=2)
    add_body(
        doc,
        "The Text Generation workspace is responsible for written campaign content. It supports use cases such as LinkedIn posts, email sequence drafting, ad hook creation, and free-form text prompts tied to the selected campaign."
    )
    add_figure(doc, "22-dashboard-text-generation.png", "Figure 3.24. Dashboard Text Generation View")

    doc.add_heading("7.6 Image Generation", level=2)
    add_body(
        doc,
        "The Image Generation workspace is used to create campaign visuals and visual variations. It also supports asset listing and creative review, making it useful for preparing branded media content."
    )
    add_figure(doc, "23-dashboard-image-generation.png", "Figure 3.25. Dashboard Image Generation View")

    doc.add_heading("7.7 Video Generation", level=2)
    add_body(
        doc,
        "The Video Generation workspace supports script writing, storyboard generation, and creator brief preparation. It extends the platform from static content production into short-form video planning."
    )
    add_figure(doc, "24-dashboard-video-generation.png", "Figure 3.26. Dashboard Video Generation View")

    doc.add_heading("7.8 Performance Analytics", level=2)
    add_body(
        doc,
        "The Performance Analytics workspace presents campaign metrics such as reach, impressions, engagement rate, clicks, and conversions. It helps the user interpret performance and supports future optimization decisions."
    )
    add_figure(doc, "25-dashboard-performance-analytics.png", "Figure 3.27. Dashboard Performance Analytics View")

    doc.add_heading("8. Reusable Components and Responsive Behavior", level=1)
    add_body(
        doc,
        "The frontend is designed using reusable components such as buttons, inputs, cards, navigation elements, layout wrappers, panels, and dialogs. This improves maintainability because changes to shared interface behavior can be made centrally and reflected across the system."
    )
    add_body(
        doc,
        "The implementation also supports responsive behavior. Layouts adapt between desktop and smaller screens through grid changes, stack-based arrangements, and mobile-specific controls such as the agent selection dropdown inside the dashboard."
    )

    doc.add_heading("9. Frontend–Backend Integration", level=1)
    add_body(
        doc,
        "The frontend communicates with the backend through API service layers. Requests are organized through dedicated service files and a shared request wrapper that handles JSON payloads, authorization headers, token refresh, and error normalization."
    )
    add_body(
        doc,
        "This integration pattern is visible in authentication, campaign loading, analytics retrieval, content generation requests, image generation requests, video generation requests, and calendar-related operations. "
        "As a result, the frontend is not an isolated static interface; it is directly connected to the backend business logic and AI-assisted features."
    )

    doc.add_heading("10. Conclusion", level=1)
    add_body(
        doc,
        "The CMO.AI frontend combines a marketing-style public interface with a modular authenticated dashboard. The implemented structure supports usability, scalability, and future expansion. "
        "By using React, TypeScript, Vite, Tailwind CSS, and reusable UI patterns, the frontend provides both a visually consistent experience and a maintainable engineering foundation for an AI-powered marketing platform."
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
